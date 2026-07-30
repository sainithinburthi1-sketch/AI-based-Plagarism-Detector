import os
import re

def parse_txt(file_path_or_stream):
    """Parses text from a .txt file path or file-like object with encoding fallbacks."""
    if isinstance(file_path_or_stream, str):
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        for enc in encodings:
            try:
                with open(file_path_or_stream, 'r', encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                raise ValueError(f"Error reading .txt file: {str(e)}")
        raise ValueError("Could not decode .txt file with supported encodings.")
    else:
        # File stream (e.g., Flask request.files stream)
        if hasattr(file_path_or_stream, 'seek'):
            file_path_or_stream.seek(0)
        content = file_path_or_stream.read()
        if hasattr(file_path_or_stream, 'seek'):
            file_path_or_stream.seek(0)
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        for enc in encodings:
            try:
                return content.decode(enc)
            except UnicodeDecodeError:
                continue
        return content.decode('utf-8', errors='ignore')

def parse_pdf(file_path_or_stream):
    """Extracts text from a PDF file using pypdf/PyPDF2."""
    extracted_text = []
    try:
        import pypdf
        reader = pypdf.PdfReader(file_path_or_stream)
        for page in reader.pages:
            t = page.extract_text()
            if t:
                extracted_text.append(t)
    except Exception as e1:
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(file_path_or_stream)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    extracted_text.append(t)
        except Exception as e2:
            raise ValueError(f"Failed to parse PDF file: {str(e1)} / {str(e2)}")

    full_text = "\n".join(extracted_text).strip()
    if not full_text:
        raise ValueError("PDF file appears to be empty or contains scanned images without extractable text.")
    return full_text

def parse_docx(file_path_or_stream):
    """Extracts text from a Word .docx file using python-docx."""
    try:
        import docx
        doc = docx.Document(file_path_or_stream)
        full_text = []
        for p in doc.paragraphs:
            if p.text.strip():
                full_text.append(p.text.strip())
        # Also extract table contents if any
        for table in doc.tables:
            for row in table.rows:
                row_txt = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_txt:
                    full_text.append(row_txt)
        return "\n".join(full_text).strip()
    except Exception as e:
        raise ValueError(f"Failed to parse .docx file: {str(e)}")

def extract_text_from_file(file_obj, filename=None):
    """
    Main entry point for extracting text from uploaded file object or file path.
    :param file_obj: file path (str) or Werkzeug FileStorage / File-like stream
    :param filename: string filename (used if file_obj is a stream)
    :return: cleaned extracted text string
    """
    if isinstance(file_obj, str):
        fname = filename or file_obj
        ext = os.path.splitext(fname)[1].lower()
        if ext == '.txt':
            return clean_text(parse_txt(file_obj))
        elif ext == '.pdf':
            return clean_text(parse_pdf(file_obj))
        elif ext in ['.docx', '.doc']:
            return clean_text(parse_docx(file_obj))
        else:
            # Fallback try plain text
            return clean_text(parse_txt(file_obj))
    else:
        fname = filename or getattr(file_obj, 'filename', '')
        ext = os.path.splitext(fname)[1].lower()
        if ext == '.txt':
            return clean_text(parse_txt(file_obj))
        elif ext == '.pdf':
            return clean_text(parse_pdf(file_obj))
        elif ext in ['.docx', '.doc']:
            return clean_text(parse_docx(file_obj))
        else:
            return clean_text(parse_txt(file_obj))

def clean_text(text):
    """Normalizes whitespace and cleans extracted text."""
    if not text:
        return ""
    # Normalize multiple line breaks and spaces
    text = re.sub(r'\r\n|\r', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()
